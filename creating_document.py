from docx import Document

document = Document()

document.add_heading('Փաստաթղթի վերնագիրը', 0)

p = document.add_paragraph('Այս փաստաթղթի սովորական տեքստի պարբերություն, ')
p.add_run('մի հատվածը՝ թավատառով, ').bold = True
p.add_run(' իսկ մյուսը՝ ')
p.add_run('շեղատառով:').italic = True

document.add_heading('Առաջին մակարդակի վերնագիրը', level=1)
document.add_paragraph('Ինչ-որ մեջբերում', style='Intense Quote')

document.add_paragraph('Չհամարակալած ցուցակի տարր',
                       style='List Bullet')
document.add_paragraph('Համարակալած ցուցակի տարր',
                       style='List Number')

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Համար'
hdr_cells[1].text = 'Անվանում'
hdr_cells[2].text = 'Քանակ'

document.save('test.docx')